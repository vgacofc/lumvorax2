#!/usr/bin/env python3
"""
Script pour générer les formats PDF, DOCX et PNG du pitch deck LUMVORAX
"""

import os
import sys
from pathlib import Path

def generate_pdf_with_markdown():
    """Génère un PDF à partir du Markdown en utilisant markdown2 + reportlab"""
    try:
        from reportlab.lib.pagesizes import letter, A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
        from reportlab.lib.enums import TA_CENTER, TA_LEFT
        import markdown2
        
        # Lire le fichier Markdown
        md_path = Path("docs/LUMVORAX_PITCH_DECK.md")
        with open(md_path, 'r', encoding='utf-8') as f:
            md_content = f.read()
        
        # Convertir Markdown en HTML
        html_content = markdown2.markdown(md_content, extras=['tables', 'fenced-code-blocks'])
        
        # Créer le PDF
        pdf_path = Path("docs/LUMVORAX_PITCH_DECK.pdf")
        doc = SimpleDocTemplate(str(pdf_path), pagesize=A4,
                                rightMargin=72, leftMargin=72,
                                topMargin=72, bottomMargin=18)
        
        # Styles
        styles = getSampleStyleSheet()
        styles.add(ParagraphStyle(name='Center', alignment=TA_CENTER))
        
        story = []
        
        # Parser simple HTML vers PDF
        lines = html_content.split('\n')
        for line in lines:
            if line.strip():
                if line.startswith('<h1>'):
                    text = line.replace('<h1>', '').replace('</h1>', '')
                    story.append(Paragraph(text, styles['Heading1']))
                    story.append(Spacer(1, 12))
                elif line.startswith('<h2>'):
                    text = line.replace('<h2>', '').replace('</h2>', '')
                    story.append(Paragraph(text, styles['Heading2']))
                    story.append(Spacer(1, 12))
                elif line.startswith('<h3>'):
                    text = line.replace('<h3>', '').replace('</h3>', '')
                    story.append(Paragraph(text, styles['Heading3']))
                    story.append(Spacer(1, 12))
                elif line.startswith('<p>'):
                    text = line.replace('<p>', '').replace('</p>', '')
                    story.append(Paragraph(text, styles['Normal']))
                    story.append(Spacer(1, 12))
                elif line == '<hr />':
                    story.append(PageBreak())
        
        doc.build(story)
        print(f"✅ PDF généré: {pdf_path}")
        return True
        
    except ImportError as e:
        print(f"❌ Bibliothèques manquantes pour PDF: {e}")
        print("   Installez: pip install reportlab markdown2")
        return False
    except Exception as e:
        print(f"❌ Erreur génération PDF: {e}")
        return False

def generate_docx():
    """Génère un DOCX à partir du Markdown"""
    try:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        # Lire le fichier Markdown
        md_path = Path("docs/LUMVORAX_PITCH_DECK.md")
        with open(md_path, 'r', encoding='utf-8') as f:
            lines = f.readlines()
        
        # Créer le document
        doc = Document()
        
        # Styles
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Arial'
        font.size = Pt(11)
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            if line.startswith('# '):
                # Titre principal
                heading = doc.add_heading(line[2:], level=1)
                heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif line.startswith('## '):
                # Titre niveau 2
                doc.add_heading(line[3:], level=2)
            elif line.startswith('### '):
                # Titre niveau 3
                doc.add_heading(line[4:], level=3)
            elif line == '---':
                # Saut de page
                doc.add_page_break()
            elif line.startswith('- ') or line.startswith('* '):
                # Liste à puces
                doc.add_paragraph(line[2:], style='List Bullet')
            elif line.startswith('| '):
                # Tableau (simplifié)
                doc.add_paragraph(line, style='Normal')
            else:
                # Paragraphe normal
                doc.add_paragraph(line)
        
        # Sauvegarder
        docx_path = Path("docs/LUMVORAX_PITCH_DECK.docx")
        doc.save(str(docx_path))
        print(f"✅ DOCX généré: {docx_path}")
        return True
        
    except ImportError as e:
        print(f"❌ Bibliothèques manquantes pour DOCX: {e}")
        print("   Installez: pip install python-docx")
        return False
    except Exception as e:
        print(f"❌ Erreur génération DOCX: {e}")
        return False

def generate_png_slides():
    """Génère des slides PNG individuelles"""
    try:
        from PIL import Image, ImageDraw, ImageFont
        
        # Lire le fichier Markdown
        md_path = Path("docs/LUMVORAX_PITCH_DECK.md")
        with open(md_path, 'r', encoding='utf-8') as f:
            content = f.read()
        
        # Séparer par slides (---)
        slides = content.split('---')
        
        # Créer le dossier de sortie
        output_dir = Path("docs/pitch_slides")
        output_dir.mkdir(exist_ok=True)
        
        # Dimensions slide (16:9)
        width, height = 1920, 1080
        bg_color = (255, 255, 255)
        text_color = (0, 0, 0)
        
        for i, slide_content in enumerate(slides, 1):
            if not slide_content.strip():
                continue
            
            # Créer l'image
            img = Image.new('RGB', (width, height), bg_color)
            draw = ImageDraw.Draw(img)
            
            # Essayer de charger une police
            try:
                title_font = ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf", 60)
                text_font = ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf", 36)
            except:
                title_font = ImageFont.load_default()
                text_font = ImageFont.load_default()
            
            # Parser le contenu
            lines = slide_content.strip().split('\n')
            y_position = 100
            
            for line in lines:
                line = line.strip()
                if not line:
                    y_position += 30
                    continue
                
                if line.startswith('# '):
                    # Titre principal
                    text = line[2:]
                    draw.text((width//2, y_position), text, fill=(0, 51, 102), 
                             font=title_font, anchor="mt")
                    y_position += 100
                elif line.startswith('## '):
                    # Sous-titre
                    text = line[3:]
                    draw.text((100, y_position), text, fill=(0, 102, 204), 
                             font=text_font)
                    y_position += 60
                else:
                    # Texte normal
                    draw.text((150, y_position), line[:80], fill=text_color, 
                             font=text_font)
                    y_position += 50
                
                if y_position > height - 100:
                    break
            
            # Sauvegarder
            output_path = output_dir / f"slide_{i:02d}.png"
            img.save(str(output_path))
            print(f"✅ Slide {i} générée: {output_path}")
        
        print(f"\n✅ {len(slides)} slides PNG générées dans {output_dir}")
        return True
        
    except ImportError as e:
        print(f"❌ Bibliothèques manquantes pour PNG: {e}")
        print("   Installez: pip install Pillow")
        return False
    except Exception as e:
        print(f"❌ Erreur génération PNG: {e}")
        return False

def main():
    """Fonction principale"""
    print("=" * 60)
    print("GÉNÉRATION DES FORMATS PITCH DECK LUMVORAX")
    print("=" * 60)
    print()
    
    results = {
        'PDF': False,
        'DOCX': False,
        'PNG': False
    }
    
    # Générer PDF
    print("📄 Génération PDF...")
    results['PDF'] = generate_pdf_with_markdown()
    print()
    
    # Générer DOCX
    print("📝 Génération DOCX...")
    results['DOCX'] = generate_docx()
    print()
    
    # Générer PNG
    print("🖼️  Génération slides PNG...")
    results['PNG'] = generate_png_slides()
    print()
    
    # Résumé
    print("=" * 60)
    print("RÉSUMÉ")
    print("=" * 60)
    for format_name, success in results.items():
        status = "✅ Succès" if success else "❌ Échec"
        print(f"{format_name}: {status}")
    
    print()
    print("📁 Fichiers disponibles:")
    print("   - docs/LUMVORAX_PITCH_DECK.md (source)")
    print("   - docs/LUMVORAX_EXECUTIVE_SUMMARY.md (résumé)")
    if results['PDF']:
        print("   - docs/LUMVORAX_PITCH_DECK.pdf")
    if results['DOCX']:
        print("   - docs/LUMVORAX_PITCH_DECK.docx")
    if results['PNG']:
        print("   - docs/pitch_slides/*.png")
    
    return all(results.values())

if __name__ == "__main__":
    success = main()
    sys.exit(0 if success else 1)

# Made with Bob
